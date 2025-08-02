# app/pipelines/docx_loader.py

import docx
import logging
import os
from pathlib import Path
from typing import List, Optional
from langchain_core.documents import Document
import asyncio
from datetime import datetime

from app.core.config import get_path_settings
from app.schemas.agent import AgentConfig, DataSource
from app.pipelines.data_policy_utils import apply_data_policies
from app.pipelines.metadata_utils import generate_base_metadata

logger = logging.getLogger(__name__)

# --- Centralized Path Configuration ---
_PATH_SETTINGS = get_path_settings()
_PROJECT_ROOT_FROM_CONFIG = _PATH_SETTINGS["PROJECT_ROOT"]
_ALLOWED_DATA_DIRS_RESOLVED = _PATH_SETTINGS["ALLOWED_DATA_DIRS"]
logger.debug(f"Loaded allowed data directories for DOCX loader from central config: {[str(d) for d in _ALLOWED_DATA_DIRS_RESOLVED]}")
# --- End Centralized Configuration ---


def _is_path_safe_and_within_allowed_dirs(input_path: str) -> Path:
    resolved_path = Path(input_path).resolve()

    is_safe = False
    for allowed_dir in _ALLOWED_DATA_DIRS_RESOLVED:
        if resolved_path.is_relative_to(allowed_dir):
            is_safe = True
            break

    if not is_safe:
        raise ValueError(f"Attempted to access path outside allowed directories: {resolved_path}")

    return resolved_path


async def load(file_path: str, agent_config: Optional[AgentConfig] = None, source: Optional[DataSource] = None) -> List[Document]:
    try:
        safe_file_path = _is_path_safe_and_within_allowed_dirs(file_path)

        if not safe_file_path.exists():
            logger.error(f"Error: DOCX file not found at {safe_file_path}")
            return []
        if not safe_file_path.is_file():
            logger.error(f"Error: Provided path '{safe_file_path}' is not a file.")
            return []
        if safe_file_path.suffix.lower() not in ['.docx']:
            logger.error(f"Error: Provided file '{safe_file_path}' is not a .docx file.")
            return []

        logger.info(f"Opening DOCX file: {safe_file_path}")

        def _open_docx_blocking():
            return docx.Document(safe_file_path)

        document = await asyncio.to_thread(_open_docx_blocking)

        full_text = "\n".join([para.text for para in document.paragraphs])

        processed_text = full_text
        document_blocked = False

        if agent_config and agent_config.data_policies:
            logger.info(f"Applying data policies to {safe_file_path.name}...")
            processed_text, document_blocked = apply_data_policies(full_text, agent_config.data_policies, policy_context="document")

        if document_blocked:
            logger.warning(f"DOCX document '{safe_file_path.name}' was completely blocked by a data policy and will not be processed.")
            return []

        if processed_text.strip():
            metadata = generate_base_metadata(
                source,
                source_context=safe_file_path.name,
                source_type="file",
            )

            metadata.update(
                {
                    "source_path": str(safe_file_path.resolve()),
                    "file_name": safe_file_path.name,
                    "file_type": safe_file_path.suffix.lower(),

                    "doc_name": safe_file_path.name,
                    "source_name": safe_file_path.name,
                    "chunk_index": 0,
                }
            )

            # reproducible, stable ID
            doc_id = safe_file_path.stem
            metadata.setdefault("original_doc_id", doc_id)

            doc = Document(page_content=processed_text, metadata=metadata, id=doc_id)
            return [doc]
        else:
            logger.warning(f"DOCX file {safe_file_path.name} contained no readable text or all content was redacted/filtered. Skipping document creation.")
            return []
    except Exception as e:
        logger.error(f"Error loading .docx file {file_path}: {e}", exc_info=True)
        return []